from __future__ import annotations

import argparse
import csv
import re
import shutil
from copy import deepcopy
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "outputs" / "PROJECT9_FIVE_METHOD_FINAL_FIGURE_TABLE_PACKAGE"
LEGENDS_DIR = OUT / "Legends"
SUPPLEMENT_DIR = OUT / "Supplement"
QC_DIR = OUT / "QC"
SUP_FIG_DIR = OUT / "Supplementary_Figures"
TABLE_WORD_DIR = OUT / "Tables" / "Word"
TABLE_SOURCE_DIR = OUT / "Tables" / "SourceData"

LEGENDS_MD = LEGENDS_DIR / "Figure_Table_Legends_FINAL.md"
LEGENDS_DOCX = LEGENDS_DIR / "Figure_Table_Legends_FINAL.docx"
SUPPLEMENT_DOCX = (
    SUPPLEMENT_DIR / "Project9_FiveMethod_Supplementary_Figures_Tables_FINAL.docx"
)

METHODS = ("GraphST", "STAGATE", "SpaGCN", "BANKSY", "SEDR")

MAIN_LEGENDS = {
    "Figure 1": (
        "Study design and five-method benchmark landscape. (a) Random-seed variation "
        "was isolated by holding the dataset, preprocessing, requested cluster number, "
        "method settings, and readout fixed, followed by separate evaluations of "
        "reference accuracy, direct partition reproducibility, iso-accuracy divergence, "
        "and marker reproducibility. (b) The benchmark comprised 12 human dorsolateral "
        "prefrontal cortex Visium sections, mouse STARmap visual cortex, human breast "
        "cancer, and five consecutive MERFISH hypothalamus/preoptic sections. (c) Complete "
        "coverage included five methods, 19 entries, and 20 seeds per method-dataset unit: "
        "95 units and 1,900 runs."
    ),
    "Figure 2": (
        "Reference-score variability and partition reproducibility. (a) Standard deviation "
        "of reference ARI over 20 seeds. (b) Median direct ARI among the 190 seed-pair "
        "partitions in each of 95 method-dataset units. (c) Reference-ARI standard deviation "
        "versus partition instability, defined as one minus median pairwise partition ARI. "
        "The 12 outlined units met the descriptive thresholds reference ARI SD <= 0.02 and "
        "partition instability >= 0.30. Across all 95 units, the descriptive Spearman rho "
        "was 0.248. (d) Method-stratified distributions of unit-level median pairwise "
        "partition ARI; points denote method-dataset units."
    ),
    "Figure 3": (
        "Near-equal benchmark scores can correspond to divergent partitions. (a) Direct "
        "partition ARI for all 6,928 seed pairs with an absolute reference-ARI difference "
        "<= 0.02. A total of 1,125 pairs (16.24%) had partition ARI < 0.50, affecting "
        "55 of 95 method-dataset units. The horizontal reference line marks partition ARI "
        "= 0.50. (b-d) Deterministically selected examples for GraphST on 151670 "
        "(seeds 9 and 14), STAGATE on 151507 (seeds 2 and 4), and SpaGCN on STARmap "
        "(seeds 2 and 17). Domain labels were aligned permutation-invariantly within each "
        "example; discordant assignments identify cells or spots assigned to different "
        "aligned domains between seeds."
    ),
    "Figure 4": (
        "Partition reproducibility tracks downstream marker reproducibility. (a) Top-100 "
        "marker-set Jaccard versus partition ARI for all 6,928 primary iso-accuracy pairs. "
        "This all-pair display is descriptive because pairs share constituent seeds. "
        "(b) Within-unit Spearman correlations between partition ARI and top-100 marker "
        "Jaccard. All 94 estimable correlations were positive, with median rho = 0.695; one "
        "unit was non-estimable because marker Jaccard was constant. (c) Unit-level median "
        "marker Jaccard across low, middle, and high within-unit partition-ARI tertiles. "
        "Integrated medians were 0.724, 0.770, and 0.818, and the median paired high-minus-low "
        "difference was 0.083 (one-sided paired Wilcoxon signed-rank test, W = 4,459, "
        "P = 2.31 x 10^-17). (d) Representative GraphST example on 151507 showing markers "
        "unique to either seed and shared between seeds in an aligned domain. The results "
        "describe analytical reproducibility and do not establish a causal effect of "
        "partition agreement on marker recovery."
    ),
    "Figure 5": (
        "Multi-seed consensus improves partition reproducibility. (a) Median seed-pair ARI "
        "and split-half consensus ARI for all 95 method-dataset units. All 95 units improved, "
        "and the median gain was 0.172. (b) Median single-seed reference ARI versus the "
        "complete 20-seed consensus reference ARI; the diagonal line denotes identity. "
        "(c) Split-half consensus ARI for every method-dataset unit. Consensus partitions "
        "were constructed from the 20 seed-specific partitions with an unweighted "
        "co-association matrix and average-linkage clustering at the same requested K; "
        "split-half reproducibility compared consensus partitions from seeds 1-10 and 11-20. "
        "Co-association consensus is an established mitigation strategy and improved "
        "reproducibility here, but it does not establish biological correctness."
    ),
}

SUPPLEMENTARY_LEGENDS = {
    "Supplementary Figure S1": (
        "Complete seed-wise reference-ARI distributions. The five panels show the 20 "
        "reference-ARI values for GraphST, STAGATE, SpaGCN, BANKSY, and SEDR across all "
        "19 benchmark entries in a common dataset order."
    ),
    "Supplementary Figure S2": (
        "NMI-based score and partition reproducibility. (a) Standard deviation of reference "
        "NMI across 20 seeds for all 95 method-dataset units. (b) Median direct pairwise "
        "partition NMI across the 190 seed pairs in each unit."
    ),
    "Supplementary Figure S3": (
        "Iso-accuracy threshold sensitivity. Results are shown for absolute reference-ARI "
        "difference thresholds of 0.01, 0.02, and 0.03. Panels report the eligible-pair "
        "count, pooled median pairwise partition ARI, and fraction of eligible pairs with "
        "partition ARI < 0.50, using the same definitions at each threshold."
    ),
    "Supplementary Figure S4": (
        "Additional deterministically selected spatial examples. The examples are GraphST "
        "on 151669 (seeds 2 and 3), STAGATE on 151510 (seeds 4 and 10), SpaGCN on 151672 "
        "(seeds 10 and 13), BANKSY on 151509 (seeds 11 and 13), and GraphST on MERFISH "
        "Bregma -0.04 (seeds 7 and 14). Reference annotations, aligned seed-specific "
        "partitions, and discordant assignments are shown. The selected seeds and maps are "
        "unchanged."
    ),
    "Supplementary Figure S5": (
        "Marker-reproducibility sensitivity and extreme-pair consequence analysis. "
        "(a) Top-50 marker-set Jaccard sensitivity. (b) Whole-ranking Spearman sensitivity. "
        "(c) Deterministically selected unstable and matched-stable contrasts. The final "
        "panel is a consequence comparison and is not interpreted as an average seed effect."
    ),
    "Supplementary Figure S6": (
        "Empirical five-method ranking distributions and winner certainty. For each "
        "dataset or section, the five observed 20-seed reference-ARI distributions were "
        "compared through the complete 20^5 = 3,200,000 Cartesian combinations. These are "
        "exact empirical combinations, not independent experiments. Panels report "
        "P(rank 1), the maximum P(rank 1) and most probable winner for each entry, the "
        "20-seed score distributions for the three entries with the smallest maximum "
        "P(rank 1), expected empirical rank, P(top 2), and P(top 3). Winner certainty was "
        "dataset dependent."
    ),
    "Supplementary Figure S7": (
        "Technical repeatability controls. Identical-seed repeatability controls are shown "
        "separately from the label-permutation sanity control. Identical-seed repeats "
        "included STAGATE, GraphST, BANKSY, and SEDR runs, including SEDR on 151507 and "
        "STARmap, and yielded partition ARI = 1.0 relative to their corresponding primary "
        "outputs. NMI is shown where available. The GraphST label-permutation control "
        "verifies permutation-invariant comparison of equivalent cluster labels and is not "
        "a model rerun."
    ),
    "Supplementary Figure S8": (
        "Complete five-method consensus analysis across all 95 method-dataset units. "
        "Panels compare median single-seed partition reproducibility with split-half "
        "consensus reproducibility, median single-seed reference ARI with complete 20-seed "
        "consensus reference ARI, and the distribution of consensus reproducibility gains "
        "by method. The same co-association consensus construction and requested K were "
        "used throughout."
    ),
}

TABLE_LEGENDS = {
    "Table 1": (
        "Spatial transcriptomics datasets included in the final five-method reproducibility "
        "benchmark. The 19 entries comprise 12 DLPFC sections, STARmap, HBCA1, and five "
        "MERFISH sections. Counts and reference-domain numbers are reported from the "
        "machine-readable dataset sources."
    ),
    "Supplementary Table S1": (
        "Dataset sources, accessions, and reference-annotation provenance for all 19 "
        "benchmark entries. Full accession and sample identifiers are retained in this "
        "table. The HBCA1 20-region pathology reference was manually defined from H&E and "
        "pathological features in the original SEDR study."
    ),
    "Supplementary Table S2": (
        "Software implementations and analysis settings for GraphST, STAGATE, SpaGCN, "
        "BANKSY, and SEDR. The table reports scientifically relevant input, preprocessing, "
        "graph, training, seed-propagation, clustering, and requested-K settings. Official "
        "SpaGCN refinement-induced reductions in observed K were retained as valid "
        "end-to-end stochastic outputs. SEDR used the official representation followed by "
        "official mclust_R fixed-K final clustering."
    ),
    "Supplementary Table S3": (
        "Complete reference-score and direct partition-reproducibility summaries for the "
        "95 method-dataset units. Iso-accuracy pairs have absolute reference-ARI difference "
        "<= 0.02; divergent iso-accuracy pairs have partition ARI < 0.50. NA denotes a "
        "quantity that could not be estimated."
    ),
    "Supplementary Table S4": (
        "Complete five-method empirical ranking, downstream marker-reproducibility, and "
        "consensus summaries for the 95 method-dataset units. Empirical rank probabilities "
        "are calculated from the exact Cartesian combinations of observed seed-level "
        "reference-ARI distributions. NA denotes a non-estimable quantity; the single "
        "non-estimable partition-to-marker Spearman correlation arose from constant marker "
        "Jaccard values and was not replaced with zero."
    ),
}

FORBIDDEN = (
    "frozen",
    "locked",
    "lock_",
    "candidate",
    "pre-unblinding",
    "amendment",
    "protocol hash",
    "technical gate",
    "wrapper",
    "seed19",
    "m0.19",
    "codex",
    "internal path",
    "no_go",
)


def _set_run_font(run, size: float, *, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), "Times New Roman")


def _configure_styles(doc: Document) -> None:
    styles = doc.styles
    for name, size, bold in (
        ("Normal", 10.0, False),
        ("Title", 15.0, True),
        ("Heading 1", 12.0, True),
        ("Heading 2", 10.5, True),
    ):
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{key}"), "Times New Roman")
    normal = styles["Normal"].paragraph_format
    normal.space_before = Pt(0)
    normal.space_after = Pt(6)
    normal.line_spacing = 1.08
    styles["Title"].paragraph_format.space_after = Pt(12)
    styles["Heading 1"].paragraph_format.space_before = Pt(10)
    styles["Heading 1"].paragraph_format.space_after = Pt(5)
    styles["Heading 1"].paragraph_format.keep_with_next = True


def _set_page_number(paragraph) -> None:
    # Section footers are often linked in Word. Always replace, rather than append,
    # the field content so a newly configured section cannot duplicate page numbers.
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    _set_run_font(run, 9.0)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))


def _configure_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    _set_page_number(section.footer.paragraphs[0])


def _configure_portrait(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    _configure_footer(section)


def _configure_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    _configure_footer(section)


def _copy_section_geometry(source, target) -> None:
    for attr in (
        "orientation",
        "page_width",
        "page_height",
        "top_margin",
        "bottom_margin",
        "left_margin",
        "right_margin",
        "header_distance",
        "footer_distance",
    ):
        setattr(target, attr, getattr(source, attr))
    _configure_footer(target)


def _add_legend_paragraph(doc: Document, label: str, text: str, *, size: float = 9.5):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.03
    paragraph.paragraph_format.keep_with_next = True
    lead = paragraph.add_run(f"{label}. ")
    _set_run_font(lead, size, bold=True)
    body = paragraph.add_run(text)
    _set_run_font(body, size)
    return paragraph


def _all_legends() -> list[tuple[str, str]]:
    return [
        *MAIN_LEGENDS.items(),
        *SUPPLEMENTARY_LEGENDS.items(),
        *TABLE_LEGENDS.items(),
    ]


def _write_markdown() -> None:
    lines = ["# Figure and Table Legends", ""]
    for heading, source in (
        ("Main Figures", MAIN_LEGENDS),
        ("Supplementary Figures", SUPPLEMENTARY_LEGENDS),
        ("Tables", TABLE_LEGENDS),
    ):
        lines.extend((f"## {heading}", ""))
        for label, legend in source.items():
            lines.extend((f"### {label}", "", legend, ""))
    LEGENDS_DIR.mkdir(parents=True, exist_ok=True)
    LEGENDS_MD.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def _write_legends_docx() -> None:
    doc = Document()
    _configure_styles(doc)
    _configure_portrait(doc.sections[0])
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True
    _set_run_font(title.add_run("Figure and Table Legends"), 15.0, bold=True)
    for heading, source in (
        ("Main Figures", MAIN_LEGENDS),
        ("Supplementary Figures", SUPPLEMENTARY_LEGENDS),
        ("Tables", TABLE_LEGENDS),
    ):
        h = doc.add_paragraph(style="Heading 1")
        _set_run_font(h.add_run(heading), 12.0, bold=True)
        for label, legend in source.items():
            _add_legend_paragraph(doc, label, legend, size=9.5)
    LEGENDS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(LEGENDS_DOCX)


def _find_figure(index: int) -> Path:
    path = SUP_FIG_DIR / f"FigureS{index}.png"
    if not path.exists():
        raise FileNotFoundError(path)
    return path


def _figure_orientation(path: Path) -> str:
    with Image.open(path) as image:
        ratio = image.width / image.height
    return "landscape" if ratio >= 1.32 else "portrait"


def _add_scaled_picture(doc: Document, path: Path, section) -> None:
    with Image.open(path) as image:
        ratio = image.width / image.height
    available_width = (
        section.page_width - section.left_margin - section.right_margin
    ) / 914400
    # Leave room for the legend and footer. Long S6 uses slightly more legend height.
    available_height = (
        section.page_height - section.top_margin - section.bottom_margin
    ) / 914400 - 1.42
    width = min(available_width, available_height * ratio)
    height = width / ratio
    if height > available_height:
        height = available_height
        width = height * ratio
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width), height=Inches(height))


def _table_source_docx(key: str) -> Path:
    return TABLE_WORD_DIR / f"Supplementary_Table_{key}_FINAL.docx"


def _table_source_csv(key: str) -> Path:
    return TABLE_SOURCE_DIR / f"Supplementary_Table_{key}_FINAL.csv"


def _repeat_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    trpr.append(hdr)


def _keep_row(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    trpr.append(cant_split)


def _set_cell_margins(cell, top=55, start=70, bottom=55, end=70) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcmar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table) -> None:
    tblpr = table._tbl.tblPr
    borders = tblpr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "bottom", "insideH", "start", "end", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        if edge in ("top", "bottom"):
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "8")
            node.set(qn("w:color"), "000000")
        elif edge == "insideH":
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "2")
            node.set(qn("w:color"), "D9D9D9")
        else:
            node.set(qn("w:val"), "nil")


def _fallback_editable_table(doc: Document, csv_path: Path) -> None:
    with csv_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.reader(handle))
    if len(rows) < 2:
        raise ValueError(f"Table source has no data rows: {csv_path}")
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_table_borders(table)
    _repeat_header(table.rows[0])
    for values in rows:
        row = table.rows[0] if values is rows[0] else table.add_row()
        _keep_row(row)
        for cell, value in zip(row.cells, values, strict=True):
            cell.text = value if value not in ("", "nan", "NaN") else "NA"
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    _set_run_font(run, 9.0, bold=values is rows[0])


def _append_final_editable_table(doc: Document, key: str) -> None:
    source_docx = _table_source_docx(key)
    if source_docx.exists():
        source = Document(source_docx)
        if not source.tables:
            raise ValueError(f"No editable table in {source_docx}")
        # Keep the table before the terminal body sectPr. Appending after sectPr
        # makes Word defer all copied tables until the final section.
        body = doc._body._body
        body.insert(len(body) - 1, deepcopy(source.tables[0]._tbl))
        return
    source_csv = _table_source_csv(key)
    if not source_csv.exists():
        raise FileNotFoundError(
            f"Neither final Word table nor final CSV exists for Supplementary Table {key}"
        )
    _fallback_editable_table(doc, source_csv)


def _write_supplement_docx() -> None:
    for index in range(1, 9):
        _find_figure(index)
    for key in ("S1", "S2", "S3", "S4"):
        if not (_table_source_docx(key).exists() or _table_source_csv(key).exists()):
            raise FileNotFoundError(f"Final editable/source table not found for {key}")

    doc = Document()
    _configure_styles(doc)
    # S1 starts the document; there is deliberately no separate cover page.
    first_path = _find_figure(1)
    if _figure_orientation(first_path) == "landscape":
        _configure_landscape(doc.sections[0])
    else:
        _configure_portrait(doc.sections[0])

    for index in range(1, 9):
        path = _find_figure(index)
        orientation = _figure_orientation(path)
        if index == 1:
            section = doc.sections[0]
        else:
            section = doc.add_section(WD_SECTION_START.NEW_PAGE)
            if orientation == "landscape":
                _configure_landscape(section)
            else:
                _configure_portrait(section)
        label = f"Supplementary Figure S{index}"
        _add_legend_paragraph(doc, label, SUPPLEMENTARY_LEGENDS[label], size=9.0)
        _add_scaled_picture(doc, path, section)

    for key in ("S1", "S2", "S3", "S4"):
        source_docx = _table_source_docx(key)
        section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        if source_docx.exists():
            _copy_section_geometry(Document(source_docx).sections[0], section)
        else:
            _configure_landscape(section)
        label = f"Supplementary Table {key}"
        _add_legend_paragraph(doc, label, TABLE_LEGENDS[label], size=9.0)
        _append_final_editable_table(doc, key)

    SUPPLEMENT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(SUPPLEMENT_DOCX)


def _assert_publication_language(paths: list[Path]) -> None:
    text = "\n".join(legend for _, legend in _all_legends()).lower()
    markdown_text = LEGENDS_MD.read_text(encoding="utf-8").lower()
    for word in FORBIDDEN:
        if word in text or word in markdown_text:
            raise AssertionError(f"Publication-language scan failed: {word}")
    # Structural checks for the two Word files.
    legends_doc = Document(LEGENDS_DOCX)
    legends_text = "\n".join(p.text for p in legends_doc.paragraphs).lower()
    supplement_doc = Document(SUPPLEMENT_DOCX)
    supplement_text = "\n".join(p.text for p in supplement_doc.paragraphs).lower()
    supplement_text += "\n" + "\n".join(
        cell.text.lower() for table in supplement_doc.tables for row in table.rows for cell in row.cells
    )
    for word in FORBIDDEN:
        if word in legends_text or word in supplement_text:
            raise AssertionError(f"Word publication-language scan failed: {word}")
    expected_labels = [label for label, _ in _all_legends()]
    for label in expected_labels:
        if label.lower() not in legends_text:
            raise AssertionError(f"Missing standalone legend: {label}")
    if len(supplement_doc.tables) != 4:
        raise AssertionError(
            f"Combined supplement must contain four editable tables, found {len(supplement_doc.tables)}"
        )
    for index in range(1, 9):
        if f"supplementary figure s{index}" not in supplement_text:
            raise AssertionError(f"Missing combined supplement figure legend S{index}")
    if len(supplement_doc.inline_shapes) != 8:
        raise AssertionError(
            f"Combined supplement must contain eight figure images, found {len(supplement_doc.inline_shapes)}"
        )
    QC_DIR.mkdir(parents=True, exist_ok=True)
    report = QC_DIR / "LEGENDS_SUPPLEMENT_STRUCTURAL_VALIDATION.md"
    report.write_text(
        "# Legends and combined supplement structural validation\n\n"
        "- Publication-language scan: PASS\n"
        "- Main figure legends: 5/5\n"
        "- Supplementary figure legends: 8/8\n"
        "- Table legends: 5/5\n"
        "- Combined supplement figure images: 8/8\n"
        "- Combined supplement editable tables: 4/4\n"
        "- Manuscript files modified: none\n",
        encoding="utf-8",
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--legends-only",
        action="store_true",
        help="Build only standalone legend files when final figures or tables are not yet available.",
    )
    args = parser.parse_args()
    LEGENDS_DIR.mkdir(parents=True, exist_ok=True)
    SUPPLEMENT_DIR.mkdir(parents=True, exist_ok=True)
    QC_DIR.mkdir(parents=True, exist_ok=True)
    _write_markdown()
    _write_legends_docx()
    if args.legends_only:
        print(LEGENDS_DOCX)
        print(LEGENDS_MD)
        return
    _write_supplement_docx()
    _assert_publication_language([LEGENDS_MD, LEGENDS_DOCX, SUPPLEMENT_DOCX])
    print(LEGENDS_DOCX)
    print(LEGENDS_MD)
    print(SUPPLEMENT_DOCX)


if __name__ == "__main__":
    main()
