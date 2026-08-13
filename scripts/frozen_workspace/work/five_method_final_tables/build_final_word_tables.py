from __future__ import annotations

import json
import math
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PACKAGE = ROOT / "outputs" / "PROJECT9_FIVE_METHOD_FINAL_FIGURE_TABLE_PACKAGE"
SOURCE = PACKAGE / "Tables" / "SourceData"
WORD = PACKAGE / "Tables" / "Word"

SPECS = {
    "Table1": ("Table 1", "Table1_FINAL.csv", "Table1_FINAL.docx"),
    "S1": ("Supplementary Table S1", "Supplementary_Table_S1_FINAL.csv", "Supplementary_Table_S1_FINAL.docx"),
    "S2": ("Supplementary Table S2", "Supplementary_Table_S2_FINAL.csv", "Supplementary_Table_S2_FINAL.docx"),
    "S3": ("Supplementary Table S3", "Supplementary_Table_S3_FINAL.csv", "Supplementary_Table_S3_FINAL.docx"),
    "S4": ("Supplementary Table S4", "Supplementary_Table_S4_FINAL.csv", "Supplementary_Table_S4_FINAL.docx"),
}
LEGENDS = {
    "Table1": "Spatial transcriptomics datasets included in the final five-method reproducibility benchmark. The 19 entries comprise 12 DLPFC sections, STARmap, HBCA1 and five consecutive MERFISH sections from one imaging-based context.",
    "S1": "Dataset sources, accessions and reference-annotation provenance. Full accession and sample identifiers are retained where appropriate. The HBCA1 reference is a manual 20-region pathology annotation based on H&E and pathological features, not a clustering output.",
    "S2": "Implementations and analysis settings for GraphST, STAGATE, SpaGCN, BANKSY and SEDR. Requested K and scientific parameters were held constant across seeds. For SpaGCN, official refinement-induced reductions in observed K were retained as valid end-to-end stochastic outputs.",
    "S3": "Complete five-method reference-score and direct partition-reproducibility summaries for 95 method-dataset units. Iso-accuracy pairs have an absolute reference-ARI difference of at most 0.02; divergent pairs have partition ARI below 0.50. Values are displayed to three decimal places; source files retain full precision. NA denotes a non-estimable quantity.",
    "S4": "Complete five-method empirical-ranking, downstream-marker and consensus summaries for 95 method-dataset units. Empirical ranks are based on exhaustive Cartesian combinations of observed 20-seed reference-ARI distributions. NA denotes the one within-unit Spearman correlation that was not estimable because marker Jaccard was constant. Values are displayed at the precisions stated in the Methods; source files retain full precision.",
}

WIDTHS = {
    "Table1": [2.1, 2.4, 1.5, 3.6, 2.0, 1.8, 7.5, 2.0],
    "S1": [2.0, 2.3, 1.3, 2.8, 2.4, 4.0, 5.6, 4.8, 1.6, 1.8, 1.5, 8.0],
    "S2": [1.8, 3.0, 2.8, 2.5, 5.2, 4.3, 5.3, 4.0, 3.5, 3.7, 4.8],
    "S3": [2.0, 1.7, 2.5, 2.1, 2.5, 2.1, 2.8, 3.0, 2.8, 2.0, 3.0, 3.0, 3.0, 3.0],
    "S4": [2.0, 1.7, 2.2, 2.3, 2.2, 1.9, 1.9, 3.3, 2.5, 2.7, 2.5, 3.0, 3.0, 2.8],
}


def set_font(run, size=9, bold=False):
    run.font.name = "Times New Roman"; run.font.size = Pt(size); run.font.bold = bold; run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr(); fonts = rpr.rFonts
    if fonts is None: fonts = OxmlElement("w:rFonts"); rpr.insert(0, fonts)
    for name in ("ascii", "hAnsi", "eastAsia", "cs"): fonts.set(qn(f"w:{name}"), "Times New Roman")


def cell_margins(cell, top=45, start=55, bottom=45, end=55):
    tcpr = cell._tc.get_or_add_tcPr(); mar = tcpr.find(qn("w:tcMar"))
    if mar is None: mar = OxmlElement("w:tcMar"); tcpr.append(mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{name}"))
        if node is None: node = OxmlElement(f"w:{name}"); mar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")


def border(cell, **edges):
    tcpr = cell._tc.get_or_add_tcPr(); borders = tcpr.find(qn("w:tcBorders"))
    if borders is None: borders = OxmlElement("w:tcBorders"); tcpr.append(borders)
    for edge, values in edges.items():
        tag = "start" if edge == "left" else "end" if edge == "right" else edge
        node = borders.find(qn(f"w:{tag}"))
        if node is None: node = OxmlElement(f"w:{tag}"); borders.append(node)
        for key, value in values.items(): node.set(qn(f"w:{key}"), str(value))


def set_width(cell, width_cm):
    cell.width = Cm(width_cm); tcpr = cell._tc.get_or_add_tcPr(); width = tcpr.find(qn("w:tcW"))
    if width is None: width = OxmlElement("w:tcW"); tcpr.append(width)
    width.set(qn("w:w"), str(int(Cm(width_cm).twips))); width.set(qn("w:type"), "dxa")


def repeat_header(row):
    trpr = row._tr.get_or_add_trPr(); node = OxmlElement("w:tblHeader"); node.set(qn("w:val"), "true"); trpr.append(node)


def keep_row(row):
    trpr = row._tr.get_or_add_trPr(); node = OxmlElement("w:cantSplit"); trpr.append(node)


def configure(doc: Document, key: str):
    normal = doc.styles["Normal"]; normal.font.name = "Times New Roman"; normal.font.size = Pt(9); normal.font.color.rgb = RGBColor(0, 0, 0)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman"); normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman"); normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    section = doc.sections[0]; section.orientation = WD_ORIENT.LANDSCAPE
    if key == "Table1": section.page_width, section.page_height = Inches(11), Inches(8.5); margin = Cm(1.15)
    else: section.page_width, section.page_height = Inches(17), Inches(11); margin = Cm(0.95)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = margin
    section.header_distance = section.footer_distance = Cm(0.65)
    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(); set_font(run, 9); begin=OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin"); instr=OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text=" PAGE "; end=OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end"); run._r.extend([begin,instr,end])


def dataframe(key: str) -> pd.DataFrame:
    _, csv_name, _ = SPECS[key]; frame = pd.read_csv(SOURCE / csv_name, keep_default_na=False, dtype=str)
    if key == "Table1":
        for col in ("Spots / cells, n", "Genes, n"): frame[col] = frame[col].map(lambda x: f"{int(float(x)):,}")
        frame["Reference domains, K"] = frame["Reference domains, K"].map(lambda x: str(int(float(x))))
    elif key == "S3":
        counts = {"Iso-accuracy pairs, n", "Divergent iso-accuracy pairs, n"}
        for col in frame.columns:
            if col in ("Dataset", "Method"): continue
            frame[col] = frame[col].map(lambda x: "NA" if x.upper()=="NA" or x=="" else str(int(float(x))) if col in counts else f"{float(x):.3f}")
    elif key == "S4":
        for col in frame.columns:
            if col in ("Dataset", "Method"): continue
            def fmt(x, c=col):
                if x.upper()=="NA" or x=="": return "NA"
                value=float(x)
                if c == "Expected empirical rank": return f"{value:.2f}"
                if c == "Median empirical rank": return str(int(value)) if value.is_integer() else f"{value:.1f}"
                return f"{value:.3f}"
            frame[col] = frame[col].map(fmt)
    return frame


def add_table(doc: Document, key: str, frame: pd.DataFrame):
    title, _, _ = SPECS[key]
    paragraph = doc.add_paragraph(); paragraph.paragraph_format.space_after = Pt(3); run = paragraph.add_run(title); set_font(run, 10, True)
    paragraph = doc.add_paragraph(); paragraph.paragraph_format.space_after = Pt(5); paragraph.paragraph_format.line_spacing = 1.0; run = paragraph.add_run(LEGENDS[key]); set_font(run, 9)
    table = doc.add_table(rows=1, cols=len(frame.columns)); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
    widths = WIDTHS[key]; total = sum(widths); grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for value in widths: node=OxmlElement("w:gridCol"); node.set(qn("w:w"), str(int(Cm(value).twips))); grid.append(node)
    tblpr=table._tbl.tblPr; tblw=tblpr.find(qn("w:tblW")); tblw.set(qn("w:w"),str(int(Cm(total).twips))); tblw.set(qn("w:type"),"dxa")
    layout=tblpr.find(qn("w:tblLayout"));
    if layout is None: layout=OxmlElement("w:tblLayout"); tblpr.append(layout)
    layout.set(qn("w:type"),"fixed")
    header = table.rows[0]; repeat_header(header); keep_row(header)
    for j, name in enumerate(frame.columns):
        cell=header.cells[j]; set_width(cell,widths[j]); cell.text=str(name); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; cell_margins(cell)
        border(cell, top={"val":"single","sz":"8","color":"000000"}, bottom={"val":"single","sz":"8","color":"000000"}, left={"val":"nil"}, right={"val":"nil"})
        for p in cell.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_before=p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
            for r in p.runs: set_font(r,9,True)
    numeric=set(frame.columns[2:] if key in ("S3","S4") else [])
    if key=="Table1": numeric={"Spots / cells, n","Genes, n","Reference domains, K"}
    for _, record in frame.iterrows():
        row=table.add_row(); keep_row(row)
        for j,(name,value) in enumerate(zip(frame.columns,record)):
            cell=row.cells[j]; set_width(cell,widths[j]); cell.text="NA" if value=="" else str(value); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; cell_margins(cell,35,50,35,50)
            border(cell,top={"val":"nil"},bottom={"val":"nil"},left={"val":"nil"},right={"val":"nil"})
            for p in cell.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.RIGHT if name in numeric else WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_before=p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
                for r in p.runs: set_font(r,9)
    for cell in table.rows[-1].cells: border(cell,top={"val":"nil"},bottom={"val":"single","sz":"8","color":"000000"},left={"val":"nil"},right={"val":"nil"})


def one(key: str):
    doc=Document(); configure(doc,key); add_table(doc,key,dataframe(key)); _,_,name=SPECS[key]; doc.save(WORD/name)


def combined():
    doc=Document(); configure(doc,"S1")
    for index,key in enumerate(("S1","S2","S3","S4")):
        if index:
            section=doc.add_section(WD_SECTION.NEW_PAGE); section.orientation=WD_ORIENT.LANDSCAPE; section.page_width,section.page_height=Inches(17),Inches(11); margin=Cm(0.95); section.top_margin=section.bottom_margin=section.left_margin=section.right_margin=margin; section.header_distance=section.footer_distance=Cm(0.65)
        add_table(doc,key,dataframe(key))
    doc.save(WORD/"Supplementary_Tables_S1-S4_FINAL.docx")


def main():
    WORD.mkdir(parents=True,exist_ok=True)
    for key in SPECS: one(key)
    combined()
    corpus="\n".join(path.read_text(encoding="utf-8",errors="ignore") for path in [])
    print(json.dumps({key:len(dataframe(key)) for key in SPECS},indent=2))


if __name__=="__main__": main()
